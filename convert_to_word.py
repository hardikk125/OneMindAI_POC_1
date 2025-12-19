#!/usr/bin/env python3
"""
Convert Markdown audit report to Word document (.docx)
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading_with_style(doc, text, level=1):
    """Add heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_from_markdown(doc, markdown_table):
    """Parse markdown table and add to document"""
    lines = markdown_table.strip().split('\n')
    if len(lines) < 3:
        return
    
    # Parse header
    header = [cell.strip() for cell in lines[0].split('|')[1:-1]]
    
    # Create table
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Light Grid Accent 1'
    
    # Add header
    header_cells = table.rows[0].cells
    for i, cell_text in enumerate(header):
        header_cells[i].text = cell_text
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add rows
    for line in lines[2:]:
        if line.strip() and '|' in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(cells) == len(header):
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cells):
                    row_cells[i].text = cell_text

def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Split content into sections
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Headings
        if line.startswith('# '):
            add_heading_with_style(doc, line[2:].strip(), level=1)
        elif line.startswith('## '):
            add_heading_with_style(doc, line[3:].strip(), level=2)
        elif line.startswith('### '):
            add_heading_with_style(doc, line[4:].strip(), level=3)
        elif line.startswith('#### '):
            add_heading_with_style(doc, line[5:].strip(), level=4)
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Add code as formatted paragraph
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'Intense Quote'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        
        # Tables
        elif line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            markdown_table = '\n'.join(table_lines)
            add_table_from_markdown(doc, markdown_table)
            i -= 1
        
        # Bold text
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip()[2:-2])
            run.font.bold = True
        
        # Lists
        elif line.strip().startswith('- '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        
        # Regular paragraphs
        elif line.strip() and not line.startswith('#') and not line.startswith('|'):
            # Skip empty lines
            if line.strip():
                doc.add_paragraph(line.strip())
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Converted: {md_file} → {docx_file}")

if __name__ == '__main__':
    md_file = r'c:\Projects\OneMindAI\docs\API_CONFIG_FINAL_AUDIT_REPORT.md'
    docx_file = r'c:\Projects\OneMindAI\docs\API_CONFIG_FINAL_AUDIT_REPORT.docx'
    
    try:
        markdown_to_docx(md_file, docx_file)
        print(f"✅ Word document created successfully!")
        print(f"📄 Location: {docx_file}")
    except Exception as e:
        print(f"❌ Error: {e}")
