#!/usr/bin/env python3
"""
Convert Markdown to DOCX
Converts HARDCODED_VALUES_FULL_AUDIT.md to DOCX format
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def read_markdown(filepath):
    """Read markdown file"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_heading(doc, text, level=1):
    """Add heading to document"""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    """Add paragraph to document"""
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def parse_table(lines, start_idx):
    """Parse markdown table and return table data and next index"""
    headers = []
    rows = []
    idx = start_idx
    
    # Parse header
    if idx < len(lines) and '|' in lines[idx]:
        header_line = lines[idx].strip()
        headers = [h.strip() for h in header_line.split('|')[1:-1]]
        idx += 1
        
        # Skip separator line
        if idx < len(lines) and '---' in lines[idx]:
            idx += 1
        
        # Parse data rows
        while idx < len(lines) and '|' in lines[idx]:
            row_line = lines[idx].strip()
            if not row_line or row_line.startswith('#'):
                break
            row = [cell.strip() for cell in row_line.split('|')[1:-1]]
            if row and any(cell for cell in row):  # Only add non-empty rows
                rows.append(row)
            idx += 1
    
    return headers, rows, idx

def add_table_to_doc(doc, headers, rows):
    """Add table to document"""
    if not headers or not rows:
        return
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = cell_text

def markdown_to_docx(md_filepath, docx_filepath):
    """Convert markdown to DOCX"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('OneMind AI - Complete Hardcoded Values Audit', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content = read_markdown(md_filepath)
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
            i += 1
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
            i += 1
        elif line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 4)
            i += 1
        
        # Handle tables
        elif '|' in line:
            headers, rows, next_idx = parse_table(lines, i)
            if headers and rows:
                add_table_to_doc(doc, headers, rows)
            i = next_idx
        
        # Handle bold text (markdown style)
        elif line.startswith('**'):
            text = line.replace('**', '')
            add_paragraph(doc, text, bold=True)
            i += 1
        
        # Handle regular paragraphs
        elif line.strip() and not line.startswith('---'):
            # Clean up markdown formatting
            text = line.replace('**', '').replace('`', '').replace('_', '')
            if text.strip():
                add_paragraph(doc, text)
            i += 1
        
        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph()
            i += 1
        
        else:
            i += 1
    
    # Save document
    doc.save(docx_filepath)
    print(f"✅ Successfully converted to: {docx_filepath}")

if __name__ == '__main__':
    md_file = r'c:\Projects\OneMindAI\docs\HARDCODED_VALUES_FULL_AUDIT.md'
    docx_file = r'c:\Projects\OneMindAI\docs\HARDCODED_VALUES_FULL_AUDIT.docx'
    
    try:
        markdown_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
